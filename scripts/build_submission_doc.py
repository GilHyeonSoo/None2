from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from PIL import Image as PILImage
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table as RLTable,
    TableStyle,
)
from paths import PAPER_FIGURES_DIR, SUBMISSION_DIR, TABLE_RESULTS_DIR

SOURCE_TXT = SUBMISSION_DIR / "UAM_Submission_Ready.txt"
OUTPUT_DOCX = SUBMISSION_DIR / "UAM_Submission_Ready.docx"
OUTPUT_PDF = SUBMISSION_DIR / "UAM_Submission_Ready.pdf"

TABLE_FILES = {
    1: TABLE_RESULTS_DIR / "table_1_overall_policy_summary.csv",
    2: TABLE_RESULTS_DIR / "table_2_phase_policy_summary.csv",
    3: TABLE_RESULTS_DIR / "table_3_latency_descriptive.csv",
    4: TABLE_RESULTS_DIR / "table_4_latency_mann_whitney.csv",
    5: TABLE_RESULTS_DIR / "table_5_top_route_improvement.csv",
}

FIGURE_FILES = {
    1: PAPER_FIGURES_DIR / "figure_1_overall_policy_comparison.png",
    2: PAPER_FIGURES_DIR / "figure_2_phase_trends.png",
    3: PAPER_FIGURES_DIR / "figure_3_condition_heatmaps.png",
    4: PAPER_FIGURES_DIR / "figure_4_top_route_interruption_reduction.png",
}

KOREAN_FONT = "/System/Library/Fonts/Supplemental/AppleGothic.ttf"
LATIN_FONT = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"


def _read_blocks() -> list[str]:
    text = SOURCE_TXT.read_text(encoding="utf-8").strip()
    return [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]


def _is_heading_level_1(block: str) -> bool:
    return bool(re.match(r"^\d+\.\s", block)) or block in {"요     약", "ABSTRACT", "References"}


def _is_heading_level_2(block: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\s", block))


def _parse_placeholder_block(block: str) -> tuple[str | None, int | None, str | None]:
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    if not lines:
        return None, None, None
    table_match = re.match(r"^\[표 (\d+) 삽입 위치\]$", lines[0])
    if table_match:
        caption = lines[1] if len(lines) > 1 else f"표 {table_match.group(1)}"
        return "table", int(table_match.group(1)), caption
    figure_match = re.match(r"^\[그림 (\d+) 삽입 위치\]$", lines[0])
    if figure_match:
        caption = lines[1] if len(lines) > 1 else f"그림 {figure_match.group(1)}"
        return "figure", int(figure_match.group(1)), caption
    return None, None, None


def _set_run_font(run, size_pt: float, bold: bool = False, align_font: str = "AppleGothic") -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), align_font)
    rfonts.set(qn("w:cs"), align_font)


def _style_paragraph(paragraph, size_pt: float = 10.5, bold: bool = False, center: bool = False) -> None:
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(4)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    _set_run_font(run, size_pt=size_pt, bold=bold)


def _df_to_docx_table(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for col_idx, column in enumerate(df.columns):
        cell = table.cell(0, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(column))
        _set_run_font(run, size_pt=8.5, bold=True)

    for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = "" if pd.isna(value) else str(value)
            run = p.add_run(text)
            _set_run_font(run, size_pt=8.0, bold=False)


def build_docx(blocks: list[str]) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    i = 0
    while i < len(blocks):
        block = blocks[i]

        if i == 0:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(block)
            _set_run_font(run, size_pt=16, bold=True)
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        if i == 1:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(block)
            _set_run_font(run, size_pt=11, bold=False)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        placeholder_type, placeholder_num, placeholder_caption = _parse_placeholder_block(block)
        if placeholder_type == "table":
            table_num = placeholder_num
            caption = placeholder_caption
            cap = document.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            _set_run_font(run, size_pt=9.5, bold=True)
            cap.paragraph_format.space_after = Pt(4)

            df = pd.read_csv(TABLE_FILES[table_num])
            _df_to_docx_table(document, df)
            document.add_paragraph()
            i += 1
            continue

        if placeholder_type == "figure":
            fig_num = placeholder_num
            caption = placeholder_caption
            path = FIGURE_FILES[fig_num]
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(path), width=Cm(15.8))

            cap = document.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            _set_run_font(run, size_pt=9.5, bold=False)
            cap.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        if _is_heading_level_1(block):
            p = document.add_paragraph()
            run = p.add_run(block)
            _set_run_font(run, size_pt=12, bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            if block in {"요     약", "ABSTRACT", "References"}:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        if _is_heading_level_2(block):
            p = document.add_paragraph()
            run = p.add_run(block)
            _set_run_font(run, size_pt=11, bold=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        p = document.add_paragraph()
        run = p.add_run(block)
        _set_run_font(run, size_pt=10.5, bold=False)
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(0.6)
        i += 1

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


def _register_pdf_fonts() -> None:
    pdfmetrics.registerFont(TTFont("AppleGothic", KOREAN_FONT))
    pdfmetrics.registerFont(TTFont("TimesNewRoman", LATIN_FONT))


def _make_pdf_styles():
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "TitleK",
        parent=styles["Title"],
        fontName="AppleGothic",
        fontSize=17,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    author = ParagraphStyle(
        "AuthorK",
        parent=styles["Normal"],
        fontName="AppleGothic",
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=14,
    )
    h1 = ParagraphStyle(
        "H1K",
        parent=styles["Heading1"],
        fontName="AppleGothic",
        fontSize=12,
        leading=16,
        spaceBefore=8,
        spaceAfter=5,
    )
    h2 = ParagraphStyle(
        "H2K",
        parent=styles["Heading2"],
        fontName="AppleGothic",
        fontSize=11,
        leading=14,
        spaceBefore=6,
        spaceAfter=4,
    )
    body = ParagraphStyle(
        "BodyK",
        parent=styles["BodyText"],
        fontName="AppleGothic",
        fontSize=9.6,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=4,
        firstLineIndent=14,
    )
    caption = ParagraphStyle(
        "CaptionK",
        parent=styles["Normal"],
        fontName="AppleGothic",
        fontSize=8.8,
        leading=11,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    return title, author, h1, h2, body, caption


def _df_to_pdf_table(df: pd.DataFrame, max_width: float) -> RLTable:
    display_df = df.copy()
    display_df = display_df.fillna("")
    data = [list(display_df.columns)] + display_df.astype(str).values.tolist()
    col_width = max_width / max(1, len(display_df.columns))
    table = RLTable(data, colWidths=[col_width] * len(display_df.columns), repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "AppleGothic"),
                ("FONTSIZE", (0, 0), (-1, 0), 7.8),
                ("FONTSIZE", (0, 1), (-1, -1), 7.2),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAEAEA")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def build_pdf(blocks: list[str]) -> None:
    _register_pdf_fonts()
    title_style, author_style, h1_style, h2_style, body_style, caption_style = _make_pdf_styles()
    doc = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
    )

    story = []
    max_width = A4[0] - doc.leftMargin - doc.rightMargin
    i = 0
    while i < len(blocks):
        block = blocks[i]
        safe_text = block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if i == 0:
            story.append(Paragraph(safe_text, title_style))
            i += 1
            continue

        if i == 1:
            story.append(Paragraph(safe_text, author_style))
            i += 1
            continue

        placeholder_type, placeholder_num, placeholder_caption = _parse_placeholder_block(block)
        if placeholder_type == "table":
            table_num = placeholder_num
            caption = placeholder_caption
            story.append(Paragraph(caption, caption_style))
            df = pd.read_csv(TABLE_FILES[table_num])
            story.append(_df_to_pdf_table(df, max_width))
            story.append(Spacer(1, 0.25 * cm))
            i += 1
            continue

        if placeholder_type == "figure":
            fig_num = placeholder_num
            caption = placeholder_caption
            image_path = FIGURE_FILES[fig_num]
            with PILImage.open(image_path) as img:
                img_w, img_h = img.size
            max_img_width = max_width * 0.92
            max_img_height = 14.5 * cm
            scale = min(max_img_width / img_w, max_img_height / img_h)
            image = RLImage(str(image_path), width=img_w * scale, height=img_h * scale)
            story.append(image)
            story.append(Spacer(1, 0.08 * cm))
            story.append(Paragraph(caption, caption_style))
            i += 1
            continue

        if _is_heading_level_1(block):
            story.append(Paragraph(safe_text, h1_style))
            i += 1
            continue

        if _is_heading_level_2(block):
            story.append(Paragraph(safe_text, h2_style))
            i += 1
            continue

        story.append(Paragraph(safe_text, body_style))
        i += 1

    doc.build(story)
    print(OUTPUT_PDF)


def main() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    blocks = _read_blocks()
    build_docx(blocks)
    build_pdf(blocks)


if __name__ == "__main__":
    main()
